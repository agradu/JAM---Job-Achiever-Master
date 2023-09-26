import docx


def find_paragraphs(document):
    found_paragraphs = []

    # Iterate through all paragraphs directly in the document body
    for p, paragraph in enumerate(document.paragraphs):
        style = paragraph.style.name
        text = paragraph.text
        parent_element = f"Document Body: Paragraph {p}"
        found_paragraphs.append((style, text, parent_element))
        paragraph.text = text + "(modified)"

    # Iterate through all tables in the document
    for table in document.tables:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                for p, paragraph in enumerate(cell.paragraphs):
                    style = paragraph.style.name
                    text = paragraph.text
                    parent_element = f"Table: Row {row._index}, Cell {i}, Paragraph {p}"
                    found_paragraphs.append((style, text, parent_element))
                    paragraph.text = text + "(modified)"

    return found_paragraphs


# Example usage
document = docx.Document("cv-template-1.docx")
paragraphs = find_paragraphs(document)

# Displaying the results
for style, text, parent_element in paragraphs:
    print(f"Style: {style}\nText: {text}\nParent Element: {parent_element}\n")

document.save("test2.docx")
