import docx
import json

document = docx.Document("docx/cv-template-2.docx")


def replace_by_style(element, info):
    # Iterate through all paragraphs directly in the document body
    for paragraph in document.paragraphs:
        if paragraph.style.name == element:
            paragraph.text = info

    # Iterate through all tables in the document
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.style.name == element:
                        paragraph.text = info


with open("json/candidate.json", "r") as file:
    data = json.load(file)

replace_by_style("Name", data["name"])
replace_by_style("Surname", data["surname"])
replace_by_style("Text Birthday", data["birthday"])
replace_by_style("Text Phone", data["phone"])
replace_by_style("Text Email", data["email"])
replace_by_style("Text Address", data["address"])
replace_by_style("Text Portfolio", data["portfolio"])

document.save(f"docx/CV-{data['name']}-{data['surname']}.docx")
