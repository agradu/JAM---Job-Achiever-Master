from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

document = Document()

lines = [
    "GIULIO ARPAIA",
    "Wathever Str. 69 12045 Berlin",
    "(willing to relocate worldwide)",
    "Linkedin: giulio-arpaia-5b19b8149",
    "Email: giulio.arpaia@gmail.com",
    "Mob: 0134567890",
]

# Adjust page margins
sections = document.sections
for section in sections:
    section.top_margin = Pt(20)
    section.bottom_margin = Pt(20)
    section.left_margin = Pt(20)
    section.right_margin = Pt(20)

# Create a table with one column and one row
table = document.add_table(rows=1, cols=1)
table.style = "Table Grid"  # Apply table grid style
table.alignment = WD_ALIGN_VERTICAL.CENTER  # Center-align the table on the page

# Set cell properties
cell = table.cell(0, 0)
cell.width = Pt(400)  # Set width of the cell
cell.vertical_alignment = (
    WD_ALIGN_VERTICAL.CENTER
)  # Center-align the content vertically

# Add a paragraph to the cell
paragraph = cell.add_paragraph()
paragraph_format = paragraph.paragraph_format
paragraph_format.alignment = (
    WD_ALIGN_VERTICAL.CENTER
)  # Center-align the text within the cell vertically

for i, line in enumerate(lines):
    run = paragraph.add_run(line)
    run.font.name = "Arial"
    if i == 0:
        run.bold = True
        run.font.size = Pt(18)
    else:
        run.font.size = Pt(12)
    paragraph.add_run().add_break()  # Add line break after each line

# Add shading
shading_element = parse_xml('<w:shd {} w:fill="{}"/>'.format(nsdecls("w"), "DCDCDC"))
cell._element.get_or_add_tcPr().append(shading_element)

# Add borders
border_element = parse_xml(
    '<w:tcBorders {} w:sz="{}" w:val="single" w:space="0" w:color="auto"/>'.format(
        nsdecls("w"), Pt(1)
    )
)
cell._element.get_or_add_tcPr().append(border_element)

document.save("test.docx")
